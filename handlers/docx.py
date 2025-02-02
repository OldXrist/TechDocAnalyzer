import logging
import sys
import os
import re

from aiogram import Bot, Router, types, F
from aiogram.types import FSInputFile
from docx import Document

from os import getenv

from docx.shared import Inches
from dotenv import load_dotenv

load_dotenv()
BOT_TOKEN = getenv("BOT_TOKEN")

bot = Bot(BOT_TOKEN)
router = Router()

logging.basicConfig(level=logging.INFO, stream=sys.stdout)
logger = logging.getLogger(__name__)


# Функция для очистки значений от знаков неравенства
def clean_value(value):
    if '-' in value:
        return value.split('-')[0].strip()
    return re.sub(r"(≤|≥|<|>|не более|не менее|до)", "", value).strip()


# Функция для обработки данных
def process_extracted_data(data):
    cleaned_data = []
    prev_max_index = None  # Stores the index of the previous "max" entry
    prev_max_key = None  # Stores the key name of the previous "max" entry

    for entry in data:
        key, value = list(entry.items())[0]
        cleaned_value = clean_value(value)  # Remove unwanted characters

        # Replace ',' with '.' only for valid numbers
        try:
            cleaned_value = cleaned_value.replace(",", ".")  # Convert to dot notation
            numeric_value = float(cleaned_value)  # Convert to float
        except ValueError:
            numeric_value = None  # Keep as string if conversion fails

        # If the key contains "max" AND "нетто", store its index and key
        if "max" in key.lower() and "нетто" in key.lower():
            prev_max_index = len(cleaned_data)  # Store index where "max" is added
            prev_max_key = key
        elif "min" in key.lower() and "нетто" in key.lower() and prev_max_index is not None and prev_max_key:
            try:
                # Ensure we correctly fetch max and min as floats
                max_value = float(cleaned_data[prev_max_index][prev_max_key].replace(",", "."))
                min_value = float(cleaned_value)

                logger.info(f'Key: {key}\nmax: {max_value}\nmin: {min_value}')

                # Compute average
                avg_value = round((max_value + min_value) / 2, 1)
                logger.info(f'Avg: {avg_value}')

                # Convert back to proper string format with ','
                avg_value_str = str(avg_value).replace(".", ",")

                # Update both "max" and "min" entries
                cleaned_data[prev_max_index][prev_max_key] = avg_value_str  # Modify existing max entry
                cleaned_value = avg_value_str  # Assign new value for min
            except ValueError:
                pass  # Ignore if values are not convertible to float
        else:
            prev_max_index = None  # Store index where "max" is added
            prev_max_key = None

        # Append the updated dictionary to cleaned_data
        cleaned_data.append({key: cleaned_value})

    return cleaned_data


from docx import Document
from docx.shared import Inches


def insert_processed_data(original_file_path, processed_data, output_file_path):
    doc = Document(original_file_path)

    for table in doc.tables:
        header_row_index = None  # Track the index of the header row

        # Find the header row dynamically
        for i, row in enumerate(table.rows):
            row_texts = [cell.text.strip() for cell in row.cells]
            if "Наименование характеристики" in row_texts and "Значение характеристики" in row_texts:
                header_row_index = i  # Mark this row as the header
                break  # Stop searching after finding the headers

        if header_row_index is None:
            continue  # Skip if no headers are found

        # Add a new column at the end
        new_column = table.add_column(Inches(1.5))  # Set width for better readability

        # Insert processed data **starting after the header row**
        for i, row in enumerate(table.rows):
            if i == header_row_index:
                row.cells[-1].text = "Обработанное значение"  # Set header for new column
            elif i > header_row_index and (i - header_row_index - 1) < len(processed_data):  # Skip header row
                key, new_value = list(processed_data[i - header_row_index - 1].items())[0]
                row.cells[-1].text = new_value  # Insert processed data

    # Save the modified document
    doc.save(output_file_path)


# Функция для поиска заголовков и извлечения данных
def extract_data_from_docx(file_path):
    doc = Document(file_path)
    extracted_data = []
    value_column_index = None
    key_column_index = None

    for table in doc.tables:
        header_row_index = None  # Track the index of the header row

        # Find the header row dynamically
        for i, row in enumerate(table.rows):
            row_texts = [cell.text.strip() for cell in row.cells]

            # Check if this row contains the table headers
            if "Наименование характеристики" in row_texts and "Значение характеристики" in row_texts:
                key_column_index = row_texts.index("Наименование характеристики")
                value_column_index = row_texts.index("Значение характеристики")
                header_row_index = i  # Mark this row as the header
                break  # Stop searching after finding the headers

        # Extract only the data rows after the header row
        if header_row_index is not None:
            for row in table.rows[header_row_index + 1:]:  # Start after the header
                if len(row.cells) > max(key_column_index, value_column_index):
                    key = row.cells[key_column_index].text.strip()
                    value = row.cells[value_column_index].text.strip()
                    if key:  # Skip empty rows
                        extracted_data.append({key: value})

    return extracted_data, value_column_index


# Функция для записи данных в .py файл
def save_data_to_python_file(data, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("extracted_data = [\n")
        for item in data:
            f.write(f"  {item},\n")
        f.write("]\n")


@router.message(F.document)
async def handle_docs(message: types.Message):
    file_id = message.document.file_id
    file_path = f"temp_{file_id}.docx"
    output_file_path = f"processed_data_{file_id}.docx"
    output_py_path = f"python_{file_id}.py"

    try:
        await bot.download(file_id, destination=file_path)
        extracted_data, value_column_index = extract_data_from_docx(file_path)

        if extracted_data:
            processed_data = process_extracted_data(extracted_data)
            insert_processed_data(file_path, processed_data, output_file_path)

            # Send the updated file back to the user
            await message.answer_document(FSInputFile(output_file_path), caption="✅ Обработанный файл готов!")
        else:
            await message.answer("⚠ Не удалось извлечь данные. Проверьте, что в файле есть таблица с нужными колонками.")

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)  # Удаляем временный файл
